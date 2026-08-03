import sys
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Constantes de Marca INGECART ---
ORANGE_HEX = "FF6A00"
BLACK_HEX = "05070B"
GRAY_METAL_HEX = "7E848E"

def hex_to_rgb(hex_str):
    """Convierte hex (#FF6A00) a tupla RGB (255, 106, 0)"""
    hex_str = hex_str.lstrip('#')
    return tuple(int(hex_str[i:i+2], 16) for i in (0, 2, 4))

def set_paragraph_spacing(paragraph, before=Pt(0), after=Pt(6)):
    """Ajusta el espaciado entre párrafos."""
    paragraph.paragraph_format.space_before = before
    paragraph.paragraph_format.space_after = after

def apply_heading_style(paragraph, level, text):
    """Aplica estilos jerárquicos limpios y corporativos."""
    if level == 1:
        run = paragraph.add_run(text)
        run.font.name = 'Montserrat'
        run.font.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, before=Pt(12), after=Pt(6))
        # Línea naranja inferior
        add_border_to_paragraph(paragraph, color=ORANGE_HEX)
    elif level == 2:
        run = paragraph.add_run(text)
        run.font.name = 'Montserrat'
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
        set_paragraph_spacing(paragraph, before=Pt(10), after=Pt(4))
        # Línea naranja sutil inferior
        add_border_to_paragraph(paragraph, color=ORANGE_HEX, width=6)
    else:  # level 3 o normal
        run = paragraph.add_run(text)
        run.font.name = 'Inter'
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*hex_to_rgb(GRAY_METAL_HEX))
        set_paragraph_spacing(paragraph, before=Pt(8), after=Pt(4))

def add_border_to_paragraph(paragraph, color="FF6A00", width=12):
    """Añade un borde inferior (línea) al párrafo."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_hyperlink(paragraph, url, text):
    """Añade un hipervínculo a un párrafo (para contacto)."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), ORANGE_HEX)
    rPr.append(color)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

def set_cell_background(cell, hex_color):
    """Establece el color de fondo de una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_font_color(cell, hex_color, bold=False):
    """Cambia el color de fuente de todo el contenido de la celda."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(*hex_to_rgb(hex_color))
            if bold:
                run.font.bold = True

def parse_html_to_docx(html_path, output_path):
    """Lee el HTML, extrae el contenido y lo escribe en un DOCX estilizado."""
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')

    doc = Document()
    
    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Buscar el contenido principal (la columna izquierda)
    main_div = soup.find('div', class_='main-content')
    if not main_div:
        # Fallback: buscar todo el cuerpo
        main_div = soup.find('body')

    if not main_div:
        print("Error: No se encontró el contenido principal.")
        sys.exit(1)

    # Procesar elementos secuencialmente
    for element in main_div.find_all(recursive=False):
        if element.name is None:
            continue

        # --- TÍTULO H1 ---
        if element.name == 'h1':
            p = doc.add_paragraph()
            # Extraer texto y eliminar el span "highlight" si existe
            clean_text = element.get_text(strip=True)
            apply_heading_style(p, 1, clean_text)
            # Añadir línea técnica (opcional)
            # p = doc.add_paragraph()
            # add_border_to_paragraph(p, color=ORANGE_HEX, width=6)

        # --- SUBTÍTULO H2 ---
        elif element.name == 'h2':
            p = doc.add_paragraph()
            clean_text = element.get_text(strip=True)
            apply_heading_style(p, 2, clean_text)

        # --- SUBTÍTULO H3 ---
        elif element.name == 'h3':
            p = doc.add_paragraph()
            clean_text = element.get_text(strip=True)
            apply_heading_style(p, 3, clean_text)

        # --- PÁRRAFO ---
        elif element.name == 'p':
            text = element.get_text(strip=True)
            # Si tiene estilos especiales (como el hero), lo ignoramos o lo tratamos como normal
            if not text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Inter'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
            set_paragraph_spacing(p, before=Pt(0), after=Pt(6))
            
            # Detectar enlaces <a> dentro del párrafo (fallback simple)
            links = element.find_all('a')
            if links:
                # Si hay enlaces, rehacemos el párrafo con ellos
                # (Simplificado: borramos el párrafo y lo reconstruimos)
                p.clear()
                for child in element.children:
                    if child.name == 'a':
                        link_text = child.get_text(strip=True)
                        href = child.get('href', '#')
                        if href.startswith('mailto:'):
                            # Insertar enlace
                            add_hyperlink(p, href, link_text)
                            # Añadir espacio después
                            p.add_run(' ')
                        else:
                            add_hyperlink(p, href, link_text)
                            p.add_run(' ')
                    elif child.name == 'strong':
                        run = p.add_run(child.get_text(strip=True))
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(*hex_to_rgb(ORANGE_HEX))
                        run.font.name = 'Inter'
                        run.font.size = Pt(11)
                    else:
                        if child.string:
                            run = p.add_run(child.string.strip())
                            run.font.name = 'Inter'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
                set_paragraph_spacing(p, before=Pt(0), after=Pt(6))

        # --- LISTAS UL/OL ---
        elif element.name in ['ul', 'ol']:
            is_ordered = (element.name == 'ol')
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number' if is_ordered else 'List Bullet')
                li_text = li.get_text(strip=True)
                # Limpiar espacios
                li_text = re.sub(r'\s+', ' ', li_text)
                run = p.add_run(li_text)
                run.font.name = 'Inter'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
                # Detectar negritas dentro del li (para destacar)
                for strong in li.find_all('strong'):
                    # Esto es más complejo, pero dejamos el texto plano por ahora
                    pass
                set_paragraph_spacing(p, before=Pt(2), after=Pt(2))

        # --- TABLAS ---
        elif element.name == 'table':
            # Convertir a tabla Word
            rows = element.find_all('tr')
            if not rows:
                continue
                
            # Determinar número de columnas
            num_cols = 0
            for row in rows:
                cols = row.find_all(['th', 'td'])
                num_cols = max(num_cols, len(cols))
            
            if num_cols == 0:
                continue

            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Autoajuste
            table.autofit = True

            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                is_header = (row.find('th') is not None)
                
                for j, cell_data in enumerate(cells):
                    if j >= num_cols:
                        break
                    word_cell = table.rows[i].cells[j]
                    # Limpiar texto
                    cell_text = cell_data.get_text(strip=True)
                    word_cell.text = cell_text
                    
                    # --- ESTILOS DE CELDA ---
                    # Encabezados
                    if is_header or i == 0:
                        set_cell_background(word_cell, "F2F2F2")
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
                                run.font.size = Pt(10)
                                run.font.name = 'Inter'
                    else:
                        # Celdas normales
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = 'Inter'
                                run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
                        
                        # Si es la columna de mejora (última) y tiene formato de mejora (+ o -)
                        if j == num_cols - 1:
                            # Verificar si es la columna de mejora (por contexto, normalmente es la última)
                            # Buscar el texto si contiene % o +
                            if any(c in cell_text for c in ['+', '%', '€']):
                                # Resaltar en naranja
                                for paragraph in word_cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(*hex_to_rgb(ORANGE_HEX))
                                        run.font.bold = True

                    # Alinear a la derecha si tiene clase 'text-right'
                    if cell_data.get('class') and 'text-right' in cell_data.get('class'):
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Si tiene clase 'highlight-cell'
                    if cell_data.get('class') and 'highlight-cell' in cell_data.get('class'):
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(*hex_to_rgb(ORANGE_HEX))
                                run.font.bold = True

            # Espacio después de la tabla
            doc.add_paragraph()

        # --- HR (línea horizontal) ---
        elif element.name == 'hr':
            p = doc.add_paragraph()
            add_border_to_paragraph(p, color=GRAY_METAL_HEX, width=6)
            set_paragraph_spacing(p, before=Pt(6), after=Pt(6))

        # --- DIVs genéricos (como el sidebar o feature cards) - Los saltamos o los tratamos como párrafos ---
        elif element.name == 'div':
            # Si contiene 'feature-card' o similar, lo convertimos a texto plano
            if element.get('class') and ('feature-card' in element.get('class') or 'sidebar' in element.get('class')):
                # Extraer texto
                text = element.get_text(strip=True)
                if text:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.name = 'Inter'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(*hex_to_rgb(BLACK_HEX))
                    set_paragraph_spacing(p, before=Pt(4), after=Pt(4))

    # Guardar
    doc.save(output_path)
    print(f"✅ Informe Word generado exitosamente: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python html_to_word_ingecart.py <ruta_del_archivo.html>")
        print("Ejemplo: python html_to_word_ingecart.py reporte_ingetrans.html")
        sys.exit(1)
    
    input_file = sys.argv[1]
    if not Path(input_file).exists():
        print(f"❌ Error: El archivo '{input_file}' no existe.")
        sys.exit(1)
    
    output_file = Path(input_file).stem + ".docx"
    parse_html_to_docx(input_file, output_file)